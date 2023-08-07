import openai
from docx import Document

openai.api_key = "MI_OPENAI_AI"  # Remplace con su  API 

audio_file_path = r"C:\Users\Documents\miAudio.mp3" # Remplace con su su ruta al audio audio

def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = openai.Audio.transcribe("whisper-1", audio_file)
    return transcription['text']



    



def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }


def abstract_summary_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "Eres una AI altamente calificada entrenada en comprensión y resumen de idiomas. Me gustaría que leyera el siguiente texto y lo resumiera en un párrafo abstracto conciso. Trate de retener los puntos más importantes, proporcionando un resumen coherente y legible que podría ayudar a una persona a comprender los puntos principales de la discusión sin necesidad de leer todo el texto. Evite detalles innecesarios o puntos tangenciales."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']



def key_points_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "Eres una AI competente con una especialidad en destilar información en puntos clave. Con base en el siguiente texto, identifique y enumere los puntos principales que se discutieron o mencionaron. Estas deben ser las ideas, hallazgos o temas más importantes que son cruciales para la esencia de la discusión. Su objetivo es proporcionar una lista que alguien pueda leer para comprender rápidamente de qué se habló."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']



def action_item_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "Eres un experto en IA en el análisis de conversaciones y la extracción de elementos de acción. Revise el texto e identifique cualquier tarea, tarea o acción que se haya acordado o mencionado como necesaria. Estas pueden ser tareas asignadas a individuos específicos o acciones generales que el grupo ha decidido tomar. Enumere estos elementos de acción de manera clara y concisa."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def sentiment_analysis(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "Como IA con experiencia en lenguaje y análisis de emociones, su tarea es analizar el sentimiento del siguiente texto. Considere el tono general de la discusión, la emoción que transmite el lenguaje utilizado y el contexto en el que se utilizan las palabras y frases. Indique si el sentimiento es generalmente positivo, negativo o neutral, y proporcione breves explicaciones para su análisis cuando sea posible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']



def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


    # audio_file_path = "Earningscall.wav"
transcription = transcribe_audio(audio_file_path)
minutes = meeting_minutes(transcription)
print(minutes)

save_as_docx(minutes, 'audioAnalizado.docx')
